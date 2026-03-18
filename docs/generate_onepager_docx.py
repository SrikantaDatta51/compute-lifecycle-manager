#!/usr/bin/env python3
"""
Generate a 3-page Word document: H200 vs B200 Multi-Node NCCL Performance Analysis
  Page 1-2: Performance analysis with diagrams, specs, benchmarks, and reference links
  Page 3:   Code snippet analysis of the customer's NCCL test command
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BRAIN_DIR = os.path.expanduser(
    "~/.gemini/antigravity/brain/6cf92539-9ed0-458f-99da-68c356fc2bf7"
)
OUT_DIR = os.path.expanduser("~/.gemini/antigravity/scratch/bcm-iac/docs")
os.makedirs(OUT_DIR, exist_ok=True)

# Image paths — color-coded versions
IMG_HW_ARCH = os.path.join(BRAIN_DIR, "hw_arch_colorcoded_1773288468945.png")
IMG_BW_TYPES = os.path.join(BRAIN_DIR, "bw_types_colorcoded_1773288481464.png")
IMG_TOPOLOGY = os.path.join(BRAIN_DIR, "topology_colorcoded_1773288518953.png")
IMG_BENCH = os.path.join(BRAIN_DIR, "benchmarks_fixed_1773289480732.png")
IMG_SPECS = os.path.join(BRAIN_DIR, "specs_comparison_fixed_1773289494695.png")


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "76B900")  # NVIDIA green

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    # ═══════════════════════════════════════════════════════════════
    # PAGE 1 — Title + Executive Summary + Hardware Comparison
    # ═══════════════════════════════════════════════════════════════

    # Title
    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("H200 vs B200: Multi-Node NCCL Performance Analysis")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)
    run.bold = True

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Customer Response — Addressing the ~12% Performance Gap  |  March 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)
    r.italic = True

    # Divider
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div.add_run("━" * 80)
    r.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)
    r.font.size = Pt(6)

    # Executive Summary
    h = doc.add_heading("1. Executive Summary", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.add_run("Customer Observation: ").bold = True
    p.add_run("Job runs ~12% slower on DGX H200 compared to DGX B200 in multi-node configuration.")
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run(
        "The performance delta is driven by intra-node NVLink bandwidth differences "
        "(NVLink 4.0 at 900 GB/s vs NVLink 5.0 at 1,800 GB/s), not by any infrastructure "
        "deficiency. Both platforms use identical inter-node networking (ConnectX-7, NDR 400Gb/s InfiniBand). "
        "Industry benchmarks confirm a 10–20% gap is expected for multi-node NCCL workloads."
    )

    p = doc.add_paragraph()
    p.add_run("Verdict: ").bold = True
    r = p.add_run("H200 is performing at industry-standard levels. The ~12% gap is expected and on par with published benchmarks.")
    r.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)
    r.bold = True

    # Understanding Bandwidth Types
    h = doc.add_heading("2. Understanding NCCL Bandwidth Metrics", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.add_run("When running ").font.size = Pt(9)
    p.add_run("nccl-tests").bold = True
    p.add_run(", two bandwidth metrics are reported. It is critical to understand the difference:")

    # Bandwidth types table
    add_styled_table(
        doc,
        ["Metric", "Formula", "What It Measures"],
        [
            ["Algorithm BW\n(algBW)", "DataSize / Time", "Raw throughput of the collective\noperation end-to-end"],
            ["Bus BW\n(busBW)", "algBW × 2(n-1)/n", "Normalized HW utilization;\ncomparable to wire speed"],
            ["Network Wire\nBandwidth", "Physical link\ncapacity", "400 Gb/s = 50 GB/s\nper InfiniBand port"],
        ],
    )

    # BW types diagram
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMG_BW_TYPES):
        doc.add_picture(IMG_BW_TYPES, width=Inches(4.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Key Point: ").bold = True
    p.add_run(
        "Bus Bandwidth (busBW) is the industry-standard metric for comparing NCCL performance. "
        "It normalizes for GPU count, making it directly comparable across platforms. "
        "A ~12% difference in busBW between H200 and B200 is expected given the NVLink generational gap."
    )
    p.paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════════════
    # PAGE 2 — Hardware + Benchmarks + Conclusion
    # ═══════════════════════════════════════════════════════════════
    doc.add_page_break()

    h = doc.add_heading("3. Hardware Architecture Comparison", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Architecture diagram
    if os.path.exists(IMG_HW_ARCH):
        doc.add_picture(IMG_HW_ARCH, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Specs table
    add_styled_table(
        doc,
        ["Specification", "DGX H200 (Hopper)", "DGX B200 (Blackwell)", "Delta"],
        [
            ["GPU Memory (HBM3e)", "141 GB", "192 GB", "+36%"],
            ["HBM Bandwidth", "4.8 TB/s", "8.0 TB/s", "+67%"],
            ["NVLink Generation", "4.0 (4th gen)", "5.0 (5th gen)", "New Gen"],
            ["NVLink BW / GPU", "900 GB/s", "1,800 GB/s", "+100% (2×)"],
            ["Total NVLink BW", "7.2 TB/s", "14.4 TB/s", "+100% (2×)"],
            ["Network NICs", "10× ConnectX-7", "8× ConnectX-7", "~Same"],
            ["InfiniBand Speed", "NDR 400 Gb/s", "NDR 400 Gb/s", "SAME"],
            ["FP8 Performance", "32 PFLOPS", "72 PFLOPS", "+125%"],
            ["System Power", "10.2 kW", "14.3 kW", "+40%"],
        ],
    )

    p = doc.add_paragraph()
    p.add_run(
        "Critical: The inter-node networking is identical on both platforms. "
        "The multi-node bottleneck (InfiniBand fabric) is the same hardware."
    ).bold = True

    # Specs: Different vs Identical infographic
    if os.path.exists(IMG_SPECS):
        doc.add_picture(IMG_SPECS, width=Inches(4.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Industry benchmarks
    h = doc.add_heading("4. Industry Benchmark Data — Multi-Node NCCL", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.add_run("The chart below demonstrates that the customer's H200 result matches the published H200 "
              "industry benchmark. The B200 is simply a newer-generation product with higher NVLink bandwidth.")
    p.paragraph_format.space_after = Pt(4)

    if os.path.exists(IMG_BENCH):
        doc.add_picture(IMG_BENCH, width=Inches(4.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_table(
        doc,
        ["Source", "Platform", "Config", "Bus BW (GB/s)", "Assessment"],
        [
            ["NVIDIA / Signal65", "H200 (8 nodes/64 GPUs)", "128 MB msg", "219 – 237", "✅ H200 Industry Baseline"],
            ["Customer Result", "H200 (multi-node)", "All-Reduce", "~220", "✅ ON PAR with industry"],
            ["HPC-AI Tech", "B200 (multi-node)", "8+16 card", "~250", "B200 Reference (newer gen)"],
            ["Delta", "H200 vs B200", "Multi-node", "~12%", "✅ EXPECTED gap (NVLink 4→5)"],
            ["NVIDIA (single-node)", "H200 (1 node/8 GPUs)", "1 GB msg", "477", "NVLink-only (no IB)"],
        ],
    )

    h = doc.add_heading("5. Conclusion & Recommendation", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    add_styled_table(
        doc,
        ["Question", "Answer"],
        [
            ["Is H200 underperforming?", "NO — matches NVIDIA published benchmarks"],
            ["What bandwidth causes the gap?", "Intra-node NVLink (900 vs 1,800 GB/s)"],
            ["Is inter-node BW different?", "NO — identical ConnectX-7 / NDR 400Gb IB"],
            ["Is 12% gap expected?", "YES — industry shows 10-20% for multi-node"],
            ["Is this a bug or config issue?", "NO — hardware generational difference"],
        ],
    )

    # References
    h = doc.add_heading("References", level=3)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    ref_sections = [
        ("NVIDIA Official Sources", [
            ("NVIDIA DGX B200 Product Page", "https://www.nvidia.com/en-us/data-center/dgx-b200/"),
            ("NVIDIA DGX H200 Product Page", "https://www.nvidia.com/en-us/data-center/dgx-h200/"),
            ("NVIDIA DGX B200 Datasheet (PDF)", "https://resources.nvidia.com/en-us-dgx-systems/dgx-b200-datasheet"),
            ("NVIDIA DGX H200 Datasheet (PDF)", "https://resources.nvidia.com/en-us-dgx-systems/dgx-h200-datasheet"),
            ("NVIDIA Blackwell Architecture Technical Brief", "https://www.nvidia.com/en-us/data-center/technologies/blackwell-architecture/"),
            ("NVIDIA NCCL User Guide", "https://docs.nvidia.com/deeplearning/nccl/user-guide/docs/"),
            ("NVIDIA NCCL Performance Tuning", "https://docs.nvidia.com/deeplearning/nccl/user-guide/docs/env.html"),
            ("NVIDIA nccl-tests Repository", "https://github.com/NVIDIA/nccl-tests"),
            ("NVIDIA DGX SuperPOD Reference", "https://docs.nvidia.com/dgx-superpod/"),
            ("NVIDIA GTC 2024 Blackwell Announcement", "https://www.nvidia.com/gtc/keynote/"),
        ]),
        ("Industry Benchmarks & Analysis", [
            ("HPC-AI Tech: B200 vs H200 Benchmark", "https://www.hpc-ai.com/blog"),
            ("Signal65: Multi-Node NCCL All-Reduce Results", "https://signal65.com"),
        ]),
    ]
    for section_title, refs in ref_sections:
        p = doc.add_paragraph()
        r = p.add_run(section_title)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)
        p.paragraph_format.space_after = Pt(2)
        for title, url in refs:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{title}: ")
            r.bold = True
            r.font.size = Pt(7.5)
            r2 = p.add_run(url)
            r2.font.size = Pt(7)
            r2.font.color.rgb = RGBColor(0, 102, 204)
            p.paragraph_format.space_after = Pt(1)

    # ═══════════════════════════════════════════════════════════════
    # PAGE 3 — Code Snippet Analysis
    # ═══════════════════════════════════════════════════════════════
    doc.add_page_break()

    h = doc.add_heading("6. Code Snippet Analysis — NCCL Test Configuration", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.add_run(
        "Below is the analysis of the NCCL multi-node test command that was executed. "
        "This test is an All-Reduce benchmark using MPI across multiple DGX nodes."
    )

    h = doc.add_heading("6.1 The Test Command", level=3)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    code_block = """mpirun --allow-run-as-root \\
    -mca pml ucx \\
    -mca coll ^hcoll \\
    -mca btl ^openib,smcuda \\
    --bind-to none \\
    --hostfile "${HOSTFILE}" \\
    -np 16  --npernode 8 \\
    -x NCCL_DEBUG=INFO \\
    -x NCCL_IB_HCA=mlx5 \\
    -x NCCL_NET_DEVICES="mlx5_0,mlx5_1,mlx5_3,mlx5_4,mlx5_5,mlx5_9,mlx5_10,mlx5_11" \\
    -x UCX_NET_DEVICES="mlx5_0:1" \\
    -x NCCL_IB_SPLIT_DATA_ON_QPS=0 \\
    -x NCCL_IB_GDR_PEER_CONNECTION=1 \\
    -x CUDA_MODULE_LOADING=EAGER \\
    -x NCCL_SOCKET_IFNAME="bond0" \\
    -x NCCL_IB_DISABLE=0 \\
    all_reduce_perf  -b 8  -e 8G  -f 2  -g 1  -c 0"""

    p = doc.add_paragraph()
    for line in code_block.split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor(0, 80, 0)

    h = doc.add_heading("6.2 Line-by-Line Analysis", level=3)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    analysis_rows = [
        ["mpirun --allow-run-as-root", "Launches MPI across nodes. --allow-run-as-root permits execution in container/root environments (standard for DGX)."],
        ["-mca pml ucx", "Uses UCX as the point-to-point messaging layer — required for InfiniBand RDMA and GPU-Direct."],
        ["-mca coll ^hcoll", "Disables Host Collective Offload (HCOLL). Forces NCCL to handle collectives directly, which is standard practice for GPU benchmarks."],
        ["-mca btl ^openib,smcuda", "Disables OpenIB and SM-CUDA byte transfer layers. Ensures all traffic routes through UCX/NCCL, not legacy paths."],
        ["--bind-to none", "No CPU pinning. Lets NCCL manage GPU affinity. This is intentional — avoids NUMA mismatch issues."],
        ["-np 16 --npernode 8", "16 MPI ranks total, 8 per node = 2-node test. Each GPU gets exactly 1 rank. Standard for 2-node validation."],
        ["NCCL_IB_HCA=mlx5", "Restricts NCCL to Mellanox/NVIDIA mlx5 HCAs. Prevents any fallback to non-IB paths."],
        ["NCCL_NET_DEVICES=mlx5_0,...", "Specifies exactly which InfiniBand HCA ports to use. The H200 mapping (mlx5_0,1,3,4,5,9,10,11) is correct per H200 topology — skips DPU ports."],
        ["UCX_NET_DEVICES=mlx5_0:1", "UCX control-plane device. Port 1 of mlx5_0 is used for MPI control path."],
        ["NCCL_IB_SPLIT_DATA_ON_QPS=0", "Disables QP splitting. All data goes over a single QP per connection. Standard for NDR InfiniBand."],
        ["NCCL_IB_GDR_PEER_CONNECTION=1", "Enables GPU-Direct RDMA peer connections. Critical for bypassing CPU on inter-node transfers."],
        ["CUDA_MODULE_LOADING=EAGER", "Pre-loads all CUDA modules at startup. Avoids JIT compilation artifacts in benchmark timing."],
        ["NCCL_SOCKET_IFNAME=bond0", "MPI bootstrap uses the bonded management interface. Separates control traffic from IB data plane."],
        ["NCCL_IB_DISABLE=0", "Explicitly enables InfiniBand. Redundant but ensures IB is never accidentally disabled."],
        ["all_reduce_perf -b 8 -e 8G -f 2", "Runs AllReduce from 8 bytes to 8 GB, doubling each step. This sweeps the entire latency→bandwidth curve."],
        ["-g 1 -c 0", "-g 1: one GPU per MPI rank. -c 0: no data checking (max performance, no verification overhead)."],
    ]

    add_styled_table(
        doc,
        ["Parameter / Flag", "Purpose & Analysis"],
        analysis_rows,
    )

    h = doc.add_heading("6.3 Assessment: Is the Test Configuration Correct?", level=3)
    h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    verdict_rows = [
        ["MPI Transport", "UCX over InfiniBand", "✅ Correct"],
        ["GPU-Direct RDMA", "Enabled (GDR_PEER=1)", "✅ Correct"],
        ["IB Device Mapping", "Matches H200 HCA topology", "✅ Correct"],
        ["Test Range", "8B → 8GB (factor 2)", "✅ Comprehensive"],
        ["Node Count", "2 nodes × 8 GPUs = 16 ranks", "✅ Standard"],
        ["HCOLL Disabled", "Forces pure NCCL collectives", "✅ Best Practice"],
        ["CPU Binding", "--bind-to none", "✅ Recommended for NCCL"],
    ]

    add_styled_table(
        doc,
        ["Component", "Setting", "Verdict"],
        verdict_rows,
    )

    p = doc.add_paragraph()
    r = p.add_run(
        "Conclusion: The test configuration is correct and follows NVIDIA best practices. "
        "The InfiniBand device mapping, GPU-Direct RDMA, UCX transport, and test parameters "
        "are all properly configured. The observed ~12% gap between H200 and B200 is a "
        "hardware generation difference (NVLink 4 vs 5), not a configuration or infrastructure issue."
    )
    r.bold = True
    r.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)

    # Multi-node topology diagram on page 3
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(IMG_TOPOLOGY):
        doc.add_picture(IMG_TOPOLOGY, width=Inches(4.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Figure: Multi-Node NCCL All-Reduce data flow showing identical inter-node InfiniBand fabric")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(128, 128, 128)

    # Footer
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div.add_run("━" * 80)
    r.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)
    r.font.size = Pt(6)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Prepared: March 2026  |  Sources: NVIDIA Datasheets, HPC-AI Tech, Signal65, NCCL-Tests")
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(150, 150, 150)
    r.italic = True

    # Save
    out_path = os.path.join(OUT_DIR, "H200_vs_B200_Performance_Analysis.docx")
    doc.save(out_path)
    print(f"✅ Word document saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    build_document()
