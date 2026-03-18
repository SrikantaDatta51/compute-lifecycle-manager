#!/usr/bin/env python3
"""
NLM Architecture Document — RSA/UML Style
Single unified document for team presentation.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

IMG = os.path.join(os.path.dirname(__file__), "images")
OUT = os.path.join(os.path.dirname(__file__), "NLM-Architecture-Document.docx")
doc = Document()

# ── Page setup ──
for sec in doc.sections:
    sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5)
    sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
    ft=sec.footer; ft.is_linked_to_previous=False
    fp=ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
    fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run()._r.append(parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'))

# ── Styles ──
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
for nm in ["Heading 1","Heading 2","Heading 3","Heading 4"]:
    try:
        h=doc.styles[nm]; h.font.name="Calibri"; h.font.color.rgb=RGBColor(0x1B,0x3A,0x5C)
    except: pass

def T(doc,h,rows,w=None):
    """Add bordered table with navy headers."""
    t=doc.add_table(rows=1+len(rows),cols=len(h)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,hh in enumerate(h):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(hh)
        r.bold=True; r.font.size=Pt(8.5); r.font.name="Calibri"; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=""; r=c.paragraphs[0].add_run(str(val))
            r.font.size=Pt(8.5); r.font.name="Calibri"
            if ri%2==1: c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7"/>'))
    if w:
        for i,ww in enumerate(w):
            for row in t.rows: row.cells[i].width=Inches(ww)

def P(doc,text,bold=False,italic=False,size=10.5,color=None,align=None):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.name="Calibri"; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=color
    if align: p.alignment=align

def I(doc,name,w=6.0):
    for f in os.listdir(IMG):
        if name in f and f.endswith(".png"):
            doc.add_picture(os.path.join(IMG,f),width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; return

# ═══════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════
for _ in range(5): doc.add_paragraph("")

tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("Node Lifecycle Management\n(NLM) Platform"); r.font.size=Pt(34); r.font.name="Calibri"; r.bold=True; r.font.color.rgb=RGBColor(0x1B,0x3A,0x5C)

doc.add_paragraph("")
tp2=doc.add_paragraph(); tp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=tp2.add_run("Architecture Document"); r2.font.size=Pt(20); r2.font.name="Calibri"; r2.font.color.rgb=RGBColor(0x4A,0x6F,0xA5)

doc.add_paragraph("")
ln=doc.add_paragraph(); ln.alignment=WD_ALIGN_PARAGRAPH.CENTER
r3=ln.add_run("━"*60); r3.font.color.rgb=RGBColor(0x4A,0x6F,0xA5); r3.font.size=Pt(10)
doc.add_paragraph("")

meta = [
    ("Version:", "2.0"),
    ("Date:", "March 13, 2026"),
    ("Author:", "Compute Platform Team"),
    ("Classification:", "Internal — Engineering"),
    ("Status:", "Approved for Implementation"),
]
for label,val in meta:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rl=p.add_run(label+" "); rl.bold=True; rl.font.size=Pt(11); rl.font.name="Calibri"
    rv=p.add_run(val); rv.font.size=Pt(11); rv.font.name="Calibri"

doc.add_page_break()

# ═══════════════════════════════════════════════
# TABLE OF CONTENTS (RSA style numbering)
# ═══════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)

toc_items = [
    ("1", "Executive Summary", ""),
    ("2", "Scope & Objectives", ""),
    ("2.1", "In Scope", ""),
    ("2.2", "Out of Scope (v1)", ""),
    ("2.3", "Assumptions & Constraints", ""),
    ("3", "System Context", ""),
    ("3.1", "Environment Topology", ""),
    ("3.2", "Node Inventory (GPU + CPU)", ""),
    ("4", "Logical Architecture", ""),
    ("4.1", "System Architecture Overview", ""),
    ("4.2", "Architecture Layers", ""),
    ("4.3", "Key Design Principles", ""),
    ("5", "Behavioral Architecture", ""),
    ("5.1", "Node State Machine (11 States)", ""),
    ("5.2", "State Definitions", ""),
    ("5.3", "State Transition Rules", ""),
    ("6", "Component Architecture", ""),
    ("6.1", "Failure Classification Engine", ""),
    ("6.2", "Cordon Ownership & Priority Arbitration", ""),
    ("6.3", "Incident Correlator", ""),
    ("6.4", "BMaaS Detection (No NPD)", ""),
    ("6.5", "Certification & Testing Pipeline", ""),
    ("7", "Deployment Architecture", ""),
    ("7.1", "Multi-AZ Deployment Topology", ""),
    ("7.2", "Central Infra Zone", ""),
    ("7.3", "NetBox Integration & Rack View", ""),
    ("8", "Interface Specifications", ""),
    ("8.1", "REST API Catalog", ""),
    ("8.2", "GitOps Artifacts", ""),
    ("8.3", "Webhook & Event Integrations", ""),
    ("9", "Operational Architecture", ""),
    ("9.1", "Partner Team Self-Serve Model", ""),
    ("9.2", "Persona-Based Views", ""),
    ("9.3", "Observability & Dashboard", ""),
    ("10", "Execution Plan", ""),
    ("10.1", "10-Day Sprint Plan", ""),
    ("10.2", "Value Metrics", ""),
    ("10.3", "Risk Assessment", ""),
]
for num, title, _ in toc_items:
    p = doc.add_paragraph()
    indent = "    " if "." in num else ""
    r = p.add_run(f"{indent}{num}  {title}")
    r.font.name = "Calibri"
    r.font.size = Pt(11 if "." not in num else 10)
    if "." not in num: r.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════
doc.add_heading("1  Executive Summary", level=1)

P(doc, "The Node Lifecycle Management (NLM) Platform provides unified lifecycle management for all compute nodes — GPU and CPU — across five environments spanning two Availability Zones. It replaces today's fragmented, multi-team, multi-system approach with a single state machine, a single cordon authority, and automated failure classification with partner-team self-serve operations.")
doc.add_paragraph("")

T(doc, ["Dimension","Current State","With NLM"], [
    ["Node state visibility","3 systems, 3 different answers","Single 11-state model, real-time API"],
    ["Failure classification","45 min manual SSH + triage","30 seconds, automated, routed to correct team"],
    ["Cordon ownership","6 teams cordon independently, weekly conflicts","P0–P4 priority arbitration, single authority"],
    ["BMaaS failure detection","None (no NPD on bare metal)","Full: DCGM + SMART + IPMI + IB cron detection"],
    ["CPU nodes","Idle, untested, wasted CapEx","In lifecycle: tested daily, available for workloads"],
    ["Customer drain notification","Compute team notifies manually","Customer-Facing team self-serves via Slack webhook"],
    ["RMA process","Compute team opens tickets manually","Auto-classified → DC Infra self-serves via API"],
    ["Compute Platform toil","~12 hours/week operational work","<2 hours/week (platform health only)"],
    ["Physical visibility","No rack-level view","NetBox rack elevation + NLM state overlay"],
    ["Fleet APIs","None","15 REST endpoints, self-serve for all partner teams"],
], [1.5,2.2,2.5])

doc.add_paragraph("")
P(doc, "Timeline: 10 calendar days from start to all 5 environments live.", bold=True)
P(doc, "Staffing: 1 engineer + AI coding agents, with partner team collaboration on Days 6–10.", bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 2. SCOPE & OBJECTIVES
# ═══════════════════════════════════════════════
doc.add_heading("2  Scope & Objectives", level=1)

doc.add_heading("2.1  In Scope", level=2)
T(doc, ["#","Capability","Description"], [
    ["S1","Unified node state machine","11-state lifecycle for GPU and CPU nodes across all backends"],
    ["S2","Failure classification","Rule-based engine classifying GPU, network, storage, power, thermal failures"],
    ["S3","Cordon ownership","Priority-based arbitration (P0–P4) with single uncordon authority"],
    ["S4","BMaaS detection","Cron-based health checks for environments without Kubernetes NPD"],
    ["S5","CPU node lifecycle","Discovery, L1 testing, certification, and assignment of idle CPU nodes"],
    ["S6","Partner self-serve","APIs + webhooks for Customer, DC Infra, Network, Storage, Security teams"],
    ["S7","NetBox integration","Node state sync to NetBox custom fields, rack elevation view"],
    ["S8","Observability","Grafana dashboards with persona-specific views, Slack + PagerDuty alerting"],
    ["S9","Backend adapters","BCM (cmsh), K8s (kubectl), bare metal (IPMI/Redfish) adapters"],
    ["S10","Fleet capacity API","Real-time GPU + CPU availability by AZ, SKU, and state"],
], [0.3,1.3,4.7])

doc.add_heading("2.2  Out of Scope (v1)", level=2)
T(doc, ["Item","Rationale","Revisit"], [
    ["ML-based failure prediction","Rule engine covers 95% at <500 nodes. Insufficient data for ML.","5000+ nodes"],
    ["Custom dashboard UI","Grafana covers all persona views. Custom UI adds maintenance.","Never"],
    ["Cross-AZ auto-rebalancing","Alert + suggest is safer than automated cross-AZ moves.","Quarter 3"],
    ["MAAS adapter","No MAAS environment currently in production.","If MAAS adopted"],
    ["Full NLM API server (v1)","CLI + cron + scripts sufficient for 10-day sprint. API in Phase 2.","Week 3"],
], [1.5,3.5,1.0])

doc.add_heading("2.3  Assumptions & Constraints", level=2)
T(doc, ["#","Assumption / Constraint"], [
    ["A1","NetBox is deployed and accessible in Infra Zone with rack and device data populated"],
    ["A2","Prometheus and Grafana are deployed in Infra Zone"],
    ["A3","Ansible controller in Infra Zone has SSH access to all environments"],
    ["A4","BCM head nodes are accessible via SSH from Infra Zone for cmsh operations"],
    ["A5","DCGM Exporter is deployed (or deployable) as systemd service on BMaaS nodes"],
    ["A6","PagerDuty and Slack integrations are available for alerting"],
    ["A7","Partner teams agree to self-serve operational model and accept API-based workflows"],
    ["C1","10-day timeline requires AI-assisted development (vibe coding)"],
    ["C2","Production deployment requires Staging validation (24h soak) first"],
], [0.3,6.0])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3. SYSTEM CONTEXT
# ═══════════════════════════════════════════════
doc.add_heading("3  System Context", level=1)

doc.add_heading("3.1  Environment Topology", level=2)
I(doc, "real_topology", 5.5)

T(doc, ["Environment","Backend(s)","Detection Method","NLM Adapter","IaC"], [
    ["AZ1 Prod","K8s + BCM","NPD + DCGM DaemonSet + cmsh","K8s + BCM adapters","Ansible + Terraform"],
    ["AZ1 Dev","K8s only","NPD + DCGM DaemonSet","K8s adapter","Terraform"],
    ["Staging","BCM + K8s + BMaaS","NPD (K8s) + cron (BMaaS)","All adapters","Ansible + Terraform"],
    ["AZ2 Prod BMaaS","BMaaS (bare metal)","DCGM systemd + smartctl + ipmitool + perfquery crons","Bare metal adapter","Ansible"],
    ["AZ2 Prod K8s","K8s only","NPD + DCGM DaemonSet","K8s adapter","Terraform"],
    ["Infra Zone","Management plane","N/A","Hosts NLM Global API","Terraform"],
], [0.9,0.8,1.5,1.0,0.8])

doc.add_heading("3.2  Node Inventory (GPU + CPU)", level=2)
T(doc, ["Environment","GPU Nodes","CPU Nodes","GPU SKU","CPU Status Today"], [
    ["AZ1 Prod","H200 + B200","Multiple idle","DGX H200 / DGX B200","Untracked, untested"],
    ["AZ1 Dev","Dev GPU","Multiple idle","Mixed","Untracked"],
    ["Staging","Mixed GPU","Multiple idle","Mixed","Untracked"],
    ["AZ2 Prod BMaaS","Bare metal GPU","Multiple idle","DGX B200","Untracked"],
    ["AZ2 Prod K8s","GPU","Multiple idle","Mixed","Untracked"],
], [0.9,0.8,0.8,1.0,1.2])

P(doc, "All CPU nodes will be brought into the NLM lifecycle: discovered, tested (CPU-specific L1 suite), certified, and made available for Infra Zone workloads.", italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. LOGICAL ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading("4  Logical Architecture", level=1)

doc.add_heading("4.1  System Architecture Overview", level=2)
I(doc, "system_architecture", 6.0)

doc.add_heading("4.2  Architecture Layers", level=2)
T(doc, ["Layer","Components","Responsibility"], [
    ["Data Plane","NPD, DCGM Exporter, SMART Monitor, IB Diagnostics, IPMI","Per-node health telemetry collection"],
    ["Control Plane","State Machine, Failure Classifier, Incident Correlator, Capacity Controller","Decision-making and state management"],
    ["Source of Truth","NetBox (physical inventory), PostgreSQL/SQLite (NLM state)","Authoritative inventory and state history"],
    ["Backend Adapters","BCM (cmsh), K8s (kubectl), Bare Metal (IPMI/Redfish)","Translate NLM operations to infrastructure commands"],
    ["Observability","Prometheus, Grafana, PagerDuty, Slack, NLM API","Monitoring, alerting, visualization, partner self-serve"],
], [1.0,2.7,2.8])

doc.add_heading("4.3  Key Design Principles", level=2)
T(doc, ["#","Principle","Rationale"], [
    ["DP1","Single Arbiter","NLM is the sole authority for node state transitions and cordon operations"],
    ["DP2","Backend Agnostic","Clean adapter interface allows adding backends without changing core logic"],
    ["DP3","Self-Serve Partners","Platform team builds NLM; partner teams self-serve domain operations via APIs"],
    ["DP4","Rules Over ML","At <500 nodes, deterministic rules are more reliable and debuggable than ML"],
    ["DP5","Customer Protection First","Customer-assigned nodes never tested/drained without explicit process"],
    ["DP6","GitOps Everything","All config, rules, dashboards, and deployments managed as code in bcm-iac"],
    ["DP7","Local Resilience","Each AZ's NLM agent operates independently during global API outages"],
], [0.3,1.3,4.7])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. BEHAVIORAL ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading("5  Behavioral Architecture", level=1)

doc.add_heading("5.1  Node State Machine (11 States)", level=2)
I(doc, "state_machine", 6.0)

doc.add_heading("5.2  State Definitions", level=2)
T(doc, ["#","State","Description","Tests?","Customer?","Timeout"], [
    ["1","Provisioning","OS imaging via PXE, BCM, or MAAS","❌","❌","4 hr → Maintenance"],
    ["2","Burn-In","Day-0 extended stress test (48 hr)","✅ Full","❌","72 hr → Maintenance"],
    ["3","Testing","Day-2 recertification (daily L1 pulse)","✅ Suite","❌","8 hr → Maintenance"],
    ["4","Certified Ready","All tests passed, available for assignment","❌","❌","—"],
    ["5","Customer Assigned","Running tenant workloads — PROTECTED","❌ Never","✅ Active","—"],
    ["6","Draining","Graceful workload migration before maintenance","❌","🔄 Migrating","2 hr → Emergency"],
    ["7","Emergency Drain","Critical HW risk, forced rapid drain","❌","🔄 Force","30 min → Power Off"],
    ["8","Scheduled Maint.","Planned maintenance window","❌","❌","—"],
    ["9","Repair","Active SW/FW remediation","❌","❌","—"],
    ["10","RMA","Awaiting hardware replacement (DC Infra owns)","❌","❌","—"],
    ["11","Decommissioned","End of life — terminal state","❌","❌","—"],
], [0.25,0.85,2.0,0.5,0.55,0.95])

doc.add_heading("5.3  State Transition Rules", level=2)
T(doc, ["From","Trigger","To","Owner","Priority","Automated?"], [
    ["Provisioning","provision_complete","Burn-In","NLM","P0","✅ Yes"],
    ["Provisioning","provision_timeout","Maintenance","NLM","P0","✅ Yes"],
    ["Burn-In","burn_in_pass","Certified Ready","NLM","P0","✅ Yes"],
    ["Burn-In","burn_in_fail","Repair","NLM","P0","✅ Yes"],
    ["Testing","l1_pass","Certified Ready","NLM","P0","✅ Yes"],
    ["Testing","l1_fail","Repair","NLM","P0","✅ Yes"],
    ["Certified Ready","customer_assign","Customer Assigned","Customer Eng","P4","Manual"],
    ["Customer Assigned","drain_request","Draining","Customer Team","P4","Needs approval"],
    ["Customer Assigned","emergency_hw","Emergency Drain","NLM","P0","✅ Auto + notify"],
    ["Draining","drain_complete","Scheduled Maint.","K8s Team","P2","✅ Yes"],
    ["Draining","drain_timeout","Emergency Drain","NLM","P0","✅ Yes"],
    ["Emergency Drain","drain_complete","Repair","NLM","P0","✅ Yes"],
    ["Scheduled Maint.","repair_start","Repair","Partner teams","P3","Self-serve"],
    ["Repair","repair_complete","Testing","Partner teams","P3","Self-serve"],
    ["Repair","rma_needed","RMA","NLM","P0","✅ Yes"],
    ["RMA","rma_complete","Provisioning","DC Infra","P3","Self-serve"],
    ["Any","decommission","Decommissioned","DC Infra","P0","Manual"],
], [0.9,0.85,0.9,0.7,0.4,0.6])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 6. COMPONENT ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading("6  Component Architecture", level=1)

doc.add_heading("6.1  Failure Classification Engine", level=2)
I(doc, "failure_classification", 5.5)
doc.add_paragraph("")

T(doc, ["#","Input Event","Classification","Confidence","Action","Routed To"], [
    ["1","Xid 79 — GPU fallen off bus","HW: GPU Fatal","95%","Auto-cordon → RMA","DC Infra + NVIDIA"],
    ["2","Xid 64 — ECC page retire failure","HW: GPU Memory","95%","Auto-cordon → RMA","DC Infra + NVIDIA"],
    ["3","Xid 48 — Double-bit ECC","HW: GPU Memory","98%","Auto-cordon → RMA","DC Infra"],
    ["4","Xid 94 — Contained ECC","HW: GPU","90%","Auto-cordon → Repair","Compute (auto-remedy)"],
    ["5","Xid 95 — Uncontained ECC","HW: GPU","98%","Emergency Drain → RMA","DC Infra"],
    ["6","ECC uncorrectable > 0","HW: Memory","98%","Auto-cordon → RMA","DC Infra"],
    ["7","ECC correctable > 1000/7d","HW: Memory (pred.)","85%","Schedule maintenance","Compute"],
    ["8","NVLink uncorrectable > 0","HW: NVSwitch","90%","Auto-cordon → RMA","DC Infra"],
    ["9","IB CRC > 1000/hr","HW: Transceiver","90%","Cordon → Repair","Network Team"],
    ["10","3+ IB errors same switch","NET: Switch failure","88%","Incident → Network","Network Team"],
    ["11","SMART reallocated > 10","HW: NVMe","92%","Auto-cordon → RMA","DC Infra"],
    ["12","PSU failed","HW: Power","95%","Auto-cordon → RMA","DC Infra"],
    ["13","GPU temp > 90°C","HW: Thermal critical","95%","Emergency Drain","DC Infra"],
    ["14","3+ nodes same rack fail","INFRA: Rack/PDU","88%","Incident → DC Ops","DC Infra"],
    ["15","CPU stress-ng fail","HW: CPU","90%","Cordon → Repair","Compute"],
], [0.25,1.2,0.8,0.5,1.2,0.8])

doc.add_heading("6.2  Cordon Ownership & Priority Arbitration", level=2)
I(doc, "cordon_ownership", 5.5)
doc.add_paragraph("")

T(doc, ["Priority","Owner","Trigger Examples","Can Be Overridden?","Uncordon Requires"], [
    ["P0","NLM Controller","GPU Xid fatal, ECC uncorrectable, PSU fail, thermal","❌ Never","Recertification pass"],
    ["P1","NLM / Security","Predictive failure, CVE patching, compliance hold","❌ Never","Patch applied + recert"],
    ["P2","K8s Platform Team","kubelet NotReady, heartbeat timeout","✅ If HW verified OK","kubelet healthy + NLM OK"],
    ["P3","Partner Teams","IB maintenance (Network), storage mount (Storage)","✅ After service restored","Service verified + NLM OK"],
    ["P4","Customer Engineering","Customer debug, workload resize request","✅ With customer approval","Customer approves"],
], [0.4,0.9,1.6,0.8,1.0])

P(doc, "Arbitration Rules: (1) Higher priority always wins. (2) Only the cordon owner can uncordon. (3) P0/P1 cordons are non-overridable. (4) All cordon/uncordon actions are audited. (5) Admin override requires P0 authority with separate audit logging.", italic=True, size=9.5)

doc.add_heading("6.3  Incident Correlator", level=2)
T(doc, ["Correlation Type","Trigger","Detection","Created Incident"], [
    ["Rack-level","3+ nodes in same rack fail within 5 min","Node → Rack mapping from NetBox","INC: Rack failure (routed to DC Infra)"],
    ["Switch-level","3+ IB errors on same leaf switch","Node → Switch-port from NetBox cable traces","INC: Switch failure (routed to Network)"],
    ["PDU-level","2+ nodes on same PDU lose power","Node → PDU from NetBox power feeds","INC: PDU failure (routed to DC Infra)"],
    ["Spine-level","3+ racks see IB errors simultaneously","Cross-rack correlation sweep","INC: Spine switch (routed to Network)"],
    ["Cooling","3+ nodes same CRAH show thermal warnings","Node → CRAH zone from NetBox","INC: Cooling event (routed to DC Infra)"],
], [0.8,1.5,1.7,1.8])

doc.add_heading("6.4  BMaaS Detection (No NPD)", level=2)
P(doc, "BMaaS environments run bare metal without Kubernetes. There is no NPD. Detection uses existing tools deployed as systemd services and cron jobs on every node:")

T(doc, ["Detection Method","What It Catches","Implementation","Already Available?"], [
    ["DCGM Exporter (systemd)","GPU Xid, ECC, temperature, NVLink, utilization","dcgm-exporter systemd service → Prometheus scrape","✅ Yes"],
    ["nvidia-smi cron (5 min)","GPU fallen off bus, driver crash, memory errors","Cron: nvidia-smi --query-gpu, parse CSV, alert","✅ Extend existing"],
    ["smartctl cron (5 min)","NVMe SMART warnings, wear level, media errors","Cron: smartctl -a -j, parse JSON, check thresholds","Easy add"],
    ["perfquery/ibdiagnet cron","IB link errors, CRC counters, port status","Cron: perfquery, check error counters","✅ Extend existing"],
    ["ipmitool sel cron (5 min)","PSU failure, fan failure, thermal events, BMC events","Cron: ipmitool sel list, parse critical entries","Easy add"],
    ["fleet-certify.sh (daily)","End-to-end GPU + network health (L1 daily test)","Existing fleet-validator, extended for NLM transitions","✅ Yes"],
    ["stress-ng + memtester (CPU)","CPU health, memory integrity for CPU-only nodes","Cron: stress-ng + memtester, report result","Easy add"],
], [1.1,1.4,2.0,0.8])

doc.add_heading("6.5  Certification & Testing Pipeline", level=2)
I(doc, "daily_testing_pipeline", 5.5)
doc.add_paragraph("")

T(doc, ["Level","Name","Suite","Duration","Schedule","Target Nodes"], [
    ["L1","Health Pulse","daily-quick (GPU) / cpu-l1 (CPU)","20 min","Daily 02:00 UTC","Unassigned GPU + CPU"],
    ["L2","Stress Test","gpu-burn / stress-ng","2 hr","Weekly Saturday","Unassigned"],
    ["L3","Fabric Validation","nccl-multinode","1 hr/pair","Weekly Saturday","GPU pairs"],
    ["L4","Full Certification","full-certification","3.5 hr","On-demand","Post-repair, post-RMA"],
    ["L5","Burn-In","burn-in-48h","48 hr","On-demand","New nodes, major RMA"],
], [0.3,0.8,1.3,0.6,0.9,1.0])

P(doc, "Customer-assigned nodes are NEVER tested. Testing runs only on nodes in certified_ready, testing, or burn_in states.", bold=True, italic=True, size=9.5)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 7. DEPLOYMENT ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading("7  Deployment Architecture", level=1)

doc.add_heading("7.1  Multi-AZ Deployment Topology", level=2)
I(doc, "multi_az_topology", 5.5)

T(doc, ["Environment","Backend","NLM Deployment","Rollout Day"], [
    ["Staging","BCM + K8s + BMaaS","Full NLM: all adapters + cron detection","Days 1–5 (validation)"],
    ["AZ1 Prod","K8s + BCM","BCM adapter + K8s adapter + cron","Day 6"],
    ["AZ1 Dev","K8s only","K8s adapter","Day 7"],
    ["AZ2 Prod BMaaS","Bare metal","Bare metal adapter + full cron suite","Day 8"],
    ["AZ2 Prod K8s","K8s only","K8s adapter","Day 9"],
    ["Infra Zone","Management","Global API + NetBox sync + dashboards","Day 10"],
], [0.9,0.8,2.5,0.8])

doc.add_heading("7.2  Central Infra Zone", level=2)
I(doc, "infra_zone_hub", 5.5)

T(doc, ["Infra Zone Component","Purpose","Technology"], [
    ["NLM Global API","Central REST API for all environments","FastAPI (Python), served on Infra Zone host"],
    ["NetBox","Physical inventory, rack views, cable traces","NetBox (already deployed)"],
    ["Prometheus / Grafana","Metrics aggregation, dashboards","Prometheus + Grafana (already deployed)"],
    ["Ansible Controller","NLM agent deployment to all environments","Ansible (bcm-iac playbooks)"],
    ["Terraform State","K8s infrastructure provisioning","Terraform (existing repos)"],
    ["PostgreSQL / SQLite","NLM state history, audit trail, certifications","SQLite (v1 single-AZ), PostgreSQL (multi-AZ)"],
], [1.2,2.2,2.2])

doc.add_heading("7.3  NetBox Integration & Rack View", level=2)
T(doc, ["NetBox Feature","NLM Integration","Consuming Persona"], [
    ["Rack Elevation View","NLM writes nlm_state + nlm_health as custom fields → rack shows node state","DC Ops, Capacity"],
    ["Device Inventory","NLM syncs: state, last_certified, firmware_bundle, tenant, cordon_owner","All teams"],
    ["Power Feeds / PDUs","Incident correlator uses PDU topology for grouping power failures","SRE, DC Infra"],
    ["Cable Traces","Incident correlator uses switch-port mapping for IB failure grouping","Network Team"],
    ["IP / Interface Mapping","Adapters use management IPs for SSH, IPMI, Redfish connectivity","NLM Controller"],
    ["Custom Fields","nlm_state, nlm_health_score, nlm_tenant, nlm_last_cert, nlm_cordon_owner","All teams"],
], [1.2,2.8,1.2])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 8. INTERFACE SPECIFICATIONS
# ═══════════════════════════════════════════════
doc.add_heading("8  Interface Specifications", level=1)

doc.add_heading("8.1  REST API Catalog", level=2)
T(doc, ["Endpoint","Method","Description","Consumer"], [
    ["/api/v1/fleet/capacity","GET","GPU+CPU count by state, AZ, SKU","VP, Capacity, Customer Eng"],
    ["/api/v1/fleet/capacity?by=sku,az","GET","Breakdown by H200/B200/CPU per AZ per state","Capacity Planning"],
    ["/api/v1/nodes","GET","List all nodes (filter: state, az, backend, type)","All teams"],
    ["/api/v1/nodes/{id}","GET","Full node record with state, health, certs, firmware","SRE, Customer Eng"],
    ["/api/v1/nodes/{id}/state","PUT","Trigger state transition (auth + priority check)","NLM agents, Partner teams"],
    ["/api/v1/nodes/{id}/cordon","POST","Request cordon with priority + reason","Partner teams (self-serve)"],
    ["/api/v1/nodes/{id}/uncordon","POST","Request uncordon (ownership check)","Partner teams (self-serve)"],
    ["/api/v1/incidents","GET","Active incidents grouped by rack/switch/PDU","SRE, DC Infra, Network"],
    ["/api/v1/certifications","GET","Latest certification per node with freshness","Fleet Automation"],
    ["/api/v1/racks/{id}/nodes","GET","Nodes in rack with NLM state overlay","DC Infra"],
    ["/api/v1/tenants/{id}/nodes","GET","Tenant's nodes, protection status","Customer Eng"],
    ["/api/v1/firmware/compliance","GET","Firmware versions vs required, patch status","Security"],
    ["/api/v1/health/network","GET","IB link health, CRC rates, Rx power","Network Team"],
    ["/api/v1/health/summary","GET","Fleet-wide health score by AZ","VP, SRE"],
    ["/api/v1/maintenance/schedule","POST","Schedule rolling maintenance window","DC Infra, Security"],
], [1.6,0.4,2.2,1.2])

doc.add_heading("8.2  GitOps Artifacts", level=2)
T(doc, ["Artifact","Repository Path","Deploys Via"], [
    ["State machine config","bcm-iac/nlm-controller/config/node-states.yml","Ansible on merge"],
    ["NLM global config","bcm-iac/nlm-controller/config/nlm.yml","Ansible on merge"],
    ["Failure classifier","bcm-iac/nlm-controller/nlm/classifier.py","CI/CD on merge"],
    ["Grafana dashboards","bcm-iac/nlm-controller/dashboards/*.json","Grafana provisioning"],
    ["Ansible playbooks","bcm-iac/playbooks/deploy-nlm-*.yml","Ansible Tower/AWX"],
    ["Terraform modules","terraform/modules/nlm-k8s/","Terraform apply"],
    ["Detection cron configs","bcm-iac/nlm-controller/crons/","Ansible on merge"],
    ["Alert rules","bcm-iac/nlm-controller/alerts/alertmanager.yml","Ansible on merge"],
    ["CPU test suite","bcm-iac/fleet-validator/tests/cpu-l1.sh","Part of fleet-certify.sh"],
], [1.3,2.7,1.0])

doc.add_heading("8.3  Webhook & Event Integrations", level=2)
T(doc, ["Event","Webhook Target","Payload","Consuming Team"], [
    ["HW failure classified","Slack #dc-infra-rma","Node ID, failure class, confidence, recommended action","DC Infra"],
    ["Customer drain needed","Slack #customer-ops","Node ID, tenant, ETR, reason","Customer-Facing Team"],
    ["Network incident","Slack #network-ops","Incident ID, affected nodes, switch, classification","Network Team"],
    ["Storage incident","Slack #storage-ops","Incident ID, node, NVMe details","Storage Team"],
    ["Critical alert (P0)","PagerDuty","Node, failure class, auto-actions taken","SRE On-Call"],
    ["Certification failure","Slack #compute-platform","Node, test, failure details","Compute Platform"],
    ["Capacity threshold","Slack #capacity-planning","AZ, current %, SKU breakdown","Capacity Team"],
], [1.1,1.1,2.0,1.1])

doc.add_page_break()

# ═══════════════════════════════════════════════
# 9. OPERATIONAL ARCHITECTURE
# ═══════════════════════════════════════════════
doc.add_heading("9  Operational Architecture", level=1)

doc.add_heading("9.1  Partner Team Self-Serve Model", level=2)
P(doc, "Design Principle: Compute Platform builds and maintains NLM. All domain-specific operations are self-served by partner teams through APIs, webhooks, and Slack workflows.", bold=True, size=10)

T(doc, ["Operation","Self-Serve Owner","NLM Automation","API / Tool"], [
    ["Customer drain notification","Customer-Facing Team","Drain event → Slack #customer-ops → team contacts tenant","Webhook → Customer team Slack"],
    ["Customer workload migration","K8s Platform Team","NLM sets draining → K8s team runs migration","K8s team uses /nodes/{id}/state"],
    ["RMA ticket creation","DC Infra Team","HW classified → auto-ticket → DC Infra Slack","Webhook → Jira/ServiceNow"],
    ["Physical hardware swap","DC Infra Team","DC Infra swaps → marks rma_complete in NLM","PUT /nodes/{id}/state (DC Infra auth)"],
    ["Network failure triage","Network Team","NET classified → routed to #network-ops","GET /incidents?type=network"],
    ["Transceiver/switch repair","Network Team","Network fixes → uncordons via API (P3)","POST /nodes/{id}/uncordon (P3 auth)"],
    ["Storage failure triage","Storage Team","STORAGE classified → routed to #storage-ops","GET /incidents?type=storage"],
    ["Security patching","Security Team","Maint window scheduled → team patches → marks done","PUT /nodes/{id}/state (Security auth)"],
    ["Firmware updates","DC Infra + Security","Rolling window orchestrated → DC Infra executes","POST /maintenance/schedule"],
    ["Capacity reporting","Capacity Team","Automated: real-time via API, no human in loop","GET /fleet/capacity (self-serve)"],
], [1.1,0.8,1.8,1.5])

doc.add_paragraph("")
P(doc, "Result: Compute Platform operational load after deployment = <2 hours/week (platform health monitoring only).", bold=True, size=10)

doc.add_heading("9.2  Persona-Based Views", level=2)
T(doc, ["Persona","Central View","Access"], [
    ["VP / Director","Fleet capacity (GPU+CPU by SKU), utilization %, SLA compliance","Grafana exec dashboard"],
    ["SRE On-Call","Active incidents, open cordons, node state map, cordon ownership","Grafana SRE dashboard + PagerDuty"],
    ["Customer Engineering","Tenant→node mapping, protection status, drain schedule, ETR","Grafana customer view"],
    ["Network Team","IB link health, switch-port mapping, Rx power, CRC trends","Grafana network panel"],
    ["DC Infra","Rack elevation + NLM state, power/cooling, RMA queue","NetBox rack view"],
    ["Security","Firmware compliance matrix, CVE patch status, audit trail","Grafana security panel"],
    ["Capacity Planning","Available vs assigned vs maintenance, per-AZ, per-SKU, trends","Grafana capacity dashboard"],
    ["Provisioning / BMaaS","Provision queue, burn-in status, boot failure rates","API: /nodes?state=provisioning"],
    ["K8s Platform","Node conditions, kubelet health, NLM↔K8s state parity","Grafana K8s panel"],
    ["Automation / CI-CD","Test results, certification history, fleet-certify status","API: /certifications"],
], [0.9,2.5,1.5])

doc.add_heading("9.3  Observability & Dashboard", level=2)
I(doc, "observability_dashboard", 5.5)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 10. EXECUTION PLAN
# ═══════════════════════════════════════════════
doc.add_heading("10  Execution Plan", level=1)

doc.add_heading("10.1  10-Day Sprint Plan", level=2)
I(doc, "task_quadrant", 5.0)

T(doc, ["Day","Target Environment","Deliverables","Validation"], [
    ["1","Infra Zone + Staging","11-state YAML config, Python state machine, BCM adapter, bare metal adapter","Unit tests pass, cmsh operations verified on Staging BCM"],
    ["2","Staging","Failure classifier (15 rules), K8s adapter, NetBox custom fields, self-serve webhooks","Classifier matches known failures from past 90 days"],
    ["3","Staging","Cordon priority model (P0–P4), partner Slack routing, CPU L1 suite (stress-ng)","P0 blocks P4 uncordon. CPU L1 runs on Staging CPU nodes"],
    ["4","Staging","BMaaS cron detection (DCGM/SMART/IPMI/IB), daily L1 GPU+CPU, start 24h soak","All crons running. L1 tests pass on unassigned nodes"],
    ["5","Staging + Infra Zone","Incident correlator, fleet capacity API, Grafana dashboards (10 views), soak validated","24h soak: zero false positives. Dashboard shows all Staging nodes"],
    ["6","AZ1 Prod","Deploy BCM + K8s adapters. Register GPU + CPU nodes. Enable partner self-serve","AZ1 Prod nodes visible in dashboard. Self-serve webhooks fire"],
    ["7","AZ1 Dev","Deploy K8s adapter. Run CPU L1 on idle nodes. Mark certified_ready","CPU nodes certified. Available in capacity API"],
    ["8","AZ2 Prod BMaaS","Deploy BM adapter + full cron suite (no NPD). Register GPU + CPU nodes","Cron detection running. DC Infra RMA webhook fires on test failure"],
    ["9","AZ2 Prod K8s","Deploy K8s adapter. CPU L1. All 5 environments live and reporting","All environments in global view. Network + Storage self-serve tested"],
    ["10","Infra Zone (global)","Global API, NetBox sync, all dashboards, partner onboarding, final validation","All 5 envs reporting. Rack view live. Partner teams onboarded"],
], [0.3,0.7,2.3,1.9])

doc.add_heading("10.2  Value Metrics", level=2)
T(doc, ["Metric","Before NLM","After NLM (Day 10)","Improvement"], [
    ["Failure classification time","45 min (manual SSH + triage)","30 seconds (auto-classify + route)","90× faster"],
    ["BMaaS failure detection","Zero (no NPD, blind)","Full: DCGM + SMART + IPMI + IB crons","0% → 100%"],
    ["CPU nodes tracked and available","0 (idle, wasted CapEx)","All CPU nodes: tested, certified, available","100% coverage"],
    ["Alerts per rack failure","8 raw PagerDuty pages","1 correlated incident","87% reduction"],
    ["Customer disruption","3–5 per quarter","0 (protected + notified + graceful drain)","100% reduction"],
    ["Capacity answer to VP","2 hours (3 systems)","10 seconds (single API, GPU + CPU)","720× faster"],
    ["Cordon conflicts","Weekly (6 teams, no arbiter)","Zero (P0–P4 single arbiter)","100% eliminated"],
    ["Compute Platform toil","~12 hours/week (manual ops)","<2 hours/week (platform only)","83% reduction"],
    ["Partner team self-serve","0 operations self-served","10 operations via API + Slack","Full self-serve"],
    ["Fleet APIs available","0","15 REST endpoints for all teams","0 → 15"],
], [1.3,1.5,1.8,0.7])

doc.add_heading("10.3  Risk Assessment", level=2)
T(doc, ["Risk","Probability","Impact","Mitigation"], [
    ["BCM cmsh API changes in BCM 12","Medium","High","Adapter pattern isolates changes to one module"],
    ["False positive classifications trigger unnecessary RMAs","Low","High","Confidence scoring + human-in-the-loop for RMA decisions"],
    ["NetBox sync drift (physical ≠ logical)","Medium","Medium","30-second sync with conflict detection + alerting"],
    ["NLM controller itself goes down","Low","Critical","HA deployment, self-monitoring, local AZ autonomy"],
    ["Partner teams bypass NLM for direct cordon","Medium","High","Periodic reconciliation detects out-of-band cordons"],
    ["Customer disrupted by false P0 emergency","Low","Critical","All P0 decisions logged, reviewed weekly, classifier tuning"],
], [2.0,0.7,0.7,3.0])

doc.add_paragraph("")
P(doc, "— End of Architecture Document —", italic=True, size=10, color=RGBColor(0x6B,0x6B,0x6B), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f"✅ Generated: {OUT}")
print(f"   Size: {os.path.getsize(OUT)/1024/1024:.1f} MB")
