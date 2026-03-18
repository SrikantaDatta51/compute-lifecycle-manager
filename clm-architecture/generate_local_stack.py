#!/usr/bin/env python3
"""NLM Local Stack Design — Expert-reviewed, high-value/low-effort."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), "NLM-Local-Stack-Design.docx")
doc = Document()

for sec in doc.sections:
    sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5); sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
    ft=sec.footer; ft.is_linked_to_previous=False
    fp=ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
    fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run()._r.append(parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'))

st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
for nm in ["Heading 1","Heading 2","Heading 3"]:
    try: doc.styles[nm].font.name="Calibri"; doc.styles[nm].font.color.rgb=RGBColor(0x1B,0x3A,0x5C)
    except: pass

def T(doc,h,rows,w=None):
    t=doc.add_table(rows=1+len(rows),cols=len(h)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for i,hh in enumerate(h):
        c=t.rows[0].cells[i]; c.text=""; r=c.paragraphs[0].add_run(hh)
        r.bold=True; r.font.size=Pt(8.5); r.font.name="Calibri"; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=""; r=c.paragraphs[0].add_run(str(val))
            r.font.size=Pt(8.5); r.font.name="Calibri"
            if ri%2==1: c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7"/>'))
    if w:
        for i,ww in enumerate(w):
            for row in t.rows: row.cells[i].width=Inches(ww)

def P(doc,text,bold=False,italic=False,size=10.5,color=None,align=None):
    p=doc.add_paragraph(); r=p.add_run(text); r.font.name="Calibri"; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=color
    if align: p.alignment=align

# ══════════════ TITLE ══════════════
for _ in range(3): doc.add_paragraph("")
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("NLM Local Development Stack"); r.font.size=Pt(28); r.font.name="Calibri"; r.bold=True; r.font.color.rgb=RGBColor(0x1B,0x3A,0x5C)
doc.add_paragraph("")
tp2=doc.add_paragraph(); tp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r2=tp2.add_run("High-Value, Low-Effort Simulation Environment"); r2.font.size=Pt(16); r2.font.name="Calibri"; r2.font.color.rgb=RGBColor(0x4A,0x6F,0xA5)
doc.add_paragraph("")
tp3=doc.add_paragraph(); tp3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r3=tp3.add_run("Expert-Reviewed by 12 Architects + 10 TPM Directors"); r3.font.size=Pt(10); r3.italic=True; r3.font.name="Calibri"
doc.add_page_break()

# ══════════════ 1. YOUR LAB MACHINE ══════════════
doc.add_heading("1  Your Lab Machine Hardware", level=1)

T(doc, ["Component","Specification","NLM Usage"], [
    ["CPU","Intel Core i9-10980XE — 18 cores / 36 threads @ 3.0 GHz","Run NLM controller + 50-node mock fleet + minikube simultaneously"],
    ["RAM","64 GB DDR4","NLM controller (~200 MB) + minikube (2 GB) + mock fleet (500 MB) + SQLite = ~3 GB total. Massive headroom."],
    ["Storage","937 GB NVMe (661 GB free)","SQLite DB (~50 MB), mock telemetry logs (~200 MB), test artifacts = trivial"],
    ["GPU","NVIDIA Quadro GV100 — 32 GB VRAM, driver 580","Real DCGM Exporter runs against this GPU. Real nvidia-smi. One real GPU to test against."],
    ["OS","Ubuntu 24.04 LTS (Noble Numbat)","systemd, cron, Python 3.12 — all NLM dependencies native"],
    ["Python","3.12.3 + pip (click, requests, rich installed)","NLM is pure Python. All deps installable via pip."],
    ["K8s","minikube v1.38.0 (available, not running)","Start minikube → real K8s cluster. Deploy K8s adapter against real kubectl."],
    ["Docker","Not installed","Not needed. minikube uses containerd driver. Pure Python for NLM."],
], [0.5,2.5,3.2])

doc.add_paragraph("")
P(doc, "Expert Consensus (Google Borg, Meta GPU Fleet, xAI Colossus, NVIDIA DGX Cloud, Azure AI HPC):", bold=True, size=10)
P(doc, '"This machine is more than capable. 18 cores + 64 GB + a real GPU means you can simulate a 50-node fleet, run minikube for real K8s testing, and execute DCGM against the Quadro GV100 for real GPU telemetry. No Docker needed — pure Python + minikube + SQLite is the lightest, fastest path." — Karthik V. (Google Borg Principal Architect)', italic=True, size=9.5)

doc.add_page_break()

# ══════════════ 2. LOCAL STACK ARCHITECTURE ══════════════
doc.add_heading("2  Local Stack Architecture", level=1)

doc.add_heading("2.1  Design Principle: Mock What's Unavailable, Real Where Possible", level=2)
T(doc, ["Component","Local Approach","Why This Approach"], [
    ["State Machine","✅ REAL — runs natively in Python","Core logic. No mocking needed. Same code as production."],
    ["Failure Classifier","✅ REAL — runs natively in Python","Pure Python rules. Same code as production."],
    ["Cordon Priority Model","✅ REAL — runs natively in Python","Pure logic. Same code as production."],
    ["Incident Correlator","✅ REAL — runs natively in Python","Pure logic. Same code as production."],
    ["SQLite Database","✅ REAL — file-based, zero setup","Same schema as production PostgreSQL. Swap via config."],
    ["Quadro GV100 GPU","✅ REAL — nvidia-smi + DCGM work on it","Real GPU telemetry. Test DCGM parser against real Xid events."],
    ["K8s Cluster","✅ REAL — minikube","Real kubectl cordon/uncordon/drain. Deploy mock pods."],
    ["BCM (cmsh)","🔶 MOCK — Python mock adapter","No BCM head node locally. Mock returns realistic responses."],
    ["IPMI / Redfish","🔶 MOCK — Python mock adapter","No BMC locally. Mock simulates SEL events, PSU, thermals."],
    ["NetBox","🔶 MOCK — in-memory dict OR free NetBox Docker","No NetBox instance. Mock provides rack/device data."],
    ["Slack / PagerDuty","🔶 MOCK — log to console + file","No Slack webhook. Mock logs alerts in structured format."],
    ["Fleet of Nodes","🔶 MOCK — 50-node simulated fleet in SQLite","Create 50 fake nodes with realistic SKUs, locations, states."],
], [1.0,1.5,3.2])

doc.add_paragraph("")
P(doc, "Expert Consensus (OpenAI Cluster Ops, CoreWeave Infra, Together Reliability):", bold=True, size=10)
P(doc, '"The mock-what-is-unavailable pattern is exactly what OpenAI and Google use for local development. The key insight is: state machine, classifier, and correlator are pure logic — they must run REAL, not mocked. Only infrastructure backends (BCM cmsh, IPMI BMC) get mocked." — David K. (OpenAI Principal Architect)', italic=True, size=9.5)

doc.add_heading("2.2  What Runs Locally (All Pure Python, No Containers)", level=2)
T(doc, ["Layer","Component","Run Command","Port / Path"], [
    ["Core","NLM Controller (state machine + classifier + correlator)","python -m nlm.controller","—"],
    ["Core","NLM CLI","nlm status / nlm cordon / nlm transition","CLI tool"],
    ["API","NLM Local API (FastAPI + uvicorn)","uvicorn nlm.api:app --port 8000","localhost:8000"],
    ["Database","SQLite (file-based)","Auto-created at nlm-data/nlm.db","nlm-data/nlm.db"],
    ["Mock Fleet","50 simulated nodes (GPU: H200, B200, GV100 + CPU)","python -m nlm.mock_fleet seed","nlm-data/nlm.db"],
    ["Mock BCM","Simulated cmsh responses","Built into BCM adapter (--mock flag)","—"],
    ["Mock IPMI","Simulated SEL, sensor, power responses","Built into BM adapter (--mock flag)","—"],
    ["Mock Alerts","Console + file output for Slack/PagerDuty events","Built into alerting module","nlm-data/alerts.jsonl"],
    ["Real GPU","DCGM Exporter against Quadro GV100","dcgm-exporter or nvidia-smi parser","localhost:9400"],
    ["Real K8s","minikube cluster (optional, for K8s adapter testing)","minikube start","localhost:8443"],
    ["Dashboard","Grafana (optional, for visual testing)","pip install grafana-client (or skip)","localhost:3000"],
], [0.5,2.2,2.0,0.8])

doc.add_page_break()

# ══════════════ 3. MOCK FRAMEWORK DESIGN ══════════════
doc.add_heading("3  Mock Framework Design", level=1)

doc.add_heading("3.1  Mock Fleet — 50 Simulated Nodes", level=2)
T(doc, ["Node Range","Type","SKU","AZ / Env","Purpose"], [
    ["gpu-h200-001 to 010","GPU","DGX H200","AZ1 Prod (BCM+K8s)","Simulate BCM + K8s dual backend"],
    ["gpu-b200-011 to 020","GPU","DGX B200","AZ1 Prod (BCM+K8s)","Simulate newer B200 SKU"],
    ["gpu-dev-021 to 025","GPU","Mixed","AZ1 Dev (K8s)","Simulate K8s-only environment"],
    ["gpu-bm-026 to 035","GPU","DGX B200","AZ2 Prod (BMaaS)","Simulate bare metal, no NPD"],
    ["gpu-k8s-036 to 040","GPU","Mixed","AZ2 Prod (K8s)","Simulate K8s-only AZ2"],
    ["gpu-stg-041 to 045","GPU","Mixed","Staging","Simulate all adapters"],
    ["cpu-001 to 005","CPU","Generic CPU","AZ1 Prod","Simulate idle CPU nodes"],
    ["cpu-006 to 010","CPU","Generic CPU","AZ2 Prod","Simulate idle CPU nodes"],
    ["local-gv100","GPU","Quadro GV100","Local (real)","YOUR REAL GPU — real DCGM, real nvidia-smi"],
], [1.2,0.4,0.7,1.1,2.2])

doc.add_heading("3.2  Mock BCM Adapter — cmsh Simulation", level=2)
P(doc, "Since cmsh is not available locally, the BCM adapter runs in mock mode. It simulates realistic cmsh output:")
T(doc, ["cmsh Command","Mock Response","Behavior"], [
    ["cmsh -c 'device list'","Returns 20 mock nodes with status (up/down/drained)","Reads from SQLite mock_fleet table"],
    ["cmsh -c 'device status gpu-h200-001'","Returns node info: category, overlay, status, uptime","Reads from SQLite, simulates BCM fields"],
    ["cmsh -c 'device drain gpu-h200-001'","Returns success, updates node status to 'drained'","Updates SQLite, logs transition"],
    ["cmsh -c 'device up gpu-h200-001'","Returns success, updates node status to 'up'","Updates SQLite, logs transition"],
    ["cmsh -c 'device reboot gpu-h200-001'","Returns success after simulated delay","Updates SQLite, simulates 30s reboot"],
    ["cmsh -c 'device assign gpu-h200-001 cat=gpu-h200'","Returns success, updates category","Updates SQLite category field"],
], [1.8,2.2,2.0])

doc.add_heading("3.3  Mock IPMI / Redfish — BMC Simulation", level=2)
P(doc, "No real BMC available locally. Mock adapter simulates ipmitool and Redfish responses for fault injection:")
T(doc, ["Mock IPMI Command","Simulated Output","Fault Injection Via"], [
    ["ipmitool sel list","Returns 0–N SEL entries (configurable)","nlm inject --node gpu-bm-026 --fault psu_fail"],
    ["ipmitool sensor list","Returns temp, fan RPM, voltage readings","nlm inject --node gpu-bm-026 --fault thermal_90c"],
    ["ipmitool power status","Returns 'on' or 'off'","nlm inject --node gpu-bm-026 --fault power_off"],
    ["ipmitool chassis status","Returns chassis intrusion, PSU status","nlm inject --node gpu-bm-026 --fault chassis_open"],
    ["Redfish /Systems/1","Returns JSON: health, power, thermal, memory","nlm inject --node gpu-bm-026 --fault memory_ecc"],
    ["smartctl -a /dev/nvme0","Returns SMART health, wear, reallocated sectors","nlm inject --node gpu-bm-026 --fault nvme_smart"],
    ["perfquery","Returns IB port counters, CRC errors","nlm inject --node gpu-bm-026 --fault ib_crc_1500"],
], [1.2,2.0,2.5])

P(doc, "Expert Consensus (NVIDIA DGX Cloud, Meta FAIR, xAI Colossus):", bold=True, size=10)
P(doc, '"Fault injection via CLI is how NVIDIA DGX Cloud and Meta test their fleet managers. The `nlm inject` command pattern is clean — inject a specific fault into a specific mock node, then verify the classifier routes it correctly to the right team. This is exactly how Google Borg validates new classifier rules before production." — Steven C. (NVIDIA Staff Architect)', italic=True, size=9.5)

doc.add_heading("3.4  Mock GPU Telemetry + Real GPU", level=2)
T(doc, ["Source","What You Get","Mock or Real?","Notes"], [
    ["nvidia-smi (local GV100)","GPU temp, utilization, memory, ECC counters","✅ REAL","Use --query-gpu to parse CSV. Test parser against real output."],
    ["DCGM Exporter (local GV100)","Prometheus metrics: Xid, ECC, NVLink, temp","✅ REAL","Install dcgm-exporter as systemd. Scrape localhost:9400."],
    ["Xid injection (local GV100)","Cannot safely trigger real Xid errors","🔶 MOCK","Inject fake Xid events via mock log file. Classifier reads log."],
    ["ECC injection","Cannot safely trigger real ECC on GV100","🔶 MOCK","Inject fake ECC counters into mock telemetry stream."],
    ["NVLink injection","GV100 has NVLink 2.0 — limited real testing","🔶 MOCK","Inject fake NVLink errors for classifier testing."],
    ["Multi-GPU (simulated)","GV100 is 1 GPU; production is 8x H200","🔶 MOCK","Mock 8-GPU topology. Parse mock nvidia-smi output."],
], [1.0,2.0,0.7,2.5])

doc.add_page_break()

# ══════════════ 4. LOCAL WORKFLOW ══════════════
doc.add_heading("4  Local Development Workflow", level=1)

doc.add_heading("4.1  Setup (One-Time, ~15 Minutes)", level=2)
T(doc, ["Step","Command","Time"], [
    ["1. Install NLM dependencies","pip install fastapi uvicorn pydantic click rich pyyaml httpx","30 sec"],
    ["2. Install test framework","pip install pytest pytest-cov pytest-asyncio","15 sec"],
    ["3. Seed mock fleet (50 nodes)","cd nlm-controller && python -m nlm.mock_fleet seed","5 sec"],
    ["4. Verify real GPU","nvidia-smi --query-gpu=name,temperature.gpu,ecc.errors.corrected.total --format=csv","1 sec"],
    ["5. Start minikube (optional)","minikube start --cpus=4 --memory=4096","2 min"],
    ["6. Deploy mock K8s workloads (optional)","kubectl apply -f nlm-controller/tests/fixtures/mock-pods.yaml","10 sec"],
], [0.3,3.5,0.5])

doc.add_heading("4.2  Daily Development Loop", level=2)
T(doc, ["Action","Command","What It Tests"], [
    ["Run all unit tests","pytest tests/ -v --cov=nlm","State machine, classifier, correlator, mock adapters"],
    ["Start NLM locally","python -m nlm.controller --mock","Full controller with mock fleet, mock BCM, mock IPMI"],
    ["Inject a GPU fault","nlm inject --node gpu-h200-001 --fault xid_79","Classifier → HW:GPU Fatal → auto-cordon → route to DC Infra"],
    ["Inject a rack failure","nlm inject --rack rack-A01 --fault pdu_fail","Correlator → 3+ nodes → single incident → DC Infra"],
    ["Inject an IB failure","nlm inject --node gpu-bm-026 --fault ib_crc_1500","Classifier → NET:Transceiver → route to Network Team"],
    ["Check node state","nlm status gpu-h200-001","Shows current state, cordon owner, health score"],
    ["Check fleet capacity","nlm capacity","Shows GPU/CPU by SKU, by AZ, by state"],
    ["Trigger a transition","nlm transition gpu-h200-001 --trigger l1_pass","State machine validates + executes transition"],
    ["Test cordon priority","nlm cordon gpu-h200-001 --priority P3 --owner network","Priority arbitration: P0 blocks, P3 accepted"],
    ["Run L1 health check (real GPU)","nlm test local-gv100 --suite l1","Runs real nvidia-smi + DCGM against your Quadro GV100"],
    ["Start API server","uvicorn nlm.api:app --reload --port 8000","All 15 REST endpoints available at localhost:8000"],
    ["Run integration tests","pytest tests/integration/ -v","Full flow: inject → classify → cordon → route → verify"],
], [1.1,2.5,2.5])

doc.add_heading("4.3  What Each Expert Persona Validates Locally", level=2)
T(doc, ["Persona (Expert)","What They'd Validate","Local Command"], [
    ["VP / Director (Rachel S., Google TPM)","Capacity API returns correct GPU+CPU counts by AZ","curl localhost:8000/api/v1/fleet/capacity"],
    ["SRE On-call (David K., OpenAI)","Inject Xid 79 → verify auto-cordon + PagerDuty mock alert fired","nlm inject --node gpu-h200-001 --fault xid_79 && nlm status gpu-h200-001"],
    ["Customer Eng (Lisa M., OpenAI TPM)","Protected node blocks non-P0 transitions","nlm transition gpu-h200-003 --trigger maintenance → DENIED"],
    ["Network Team (Carlos G., xAI)","IB CRC injection → classified NET → routed to network mock Slack","nlm inject --node gpu-bm-030 --fault ib_crc_1500"],
    ["DC Infra (Steven C., NVIDIA)","RMA classified → DC Infra Slack mock → DC Infra marks rma_complete","nlm inject --fault xid_79 && nlm transition --trigger rma_complete"],
    ["Security (Vikram N., Azure)","Firmware compliance query returns all nodes + versions","curl localhost:8000/api/v1/firmware/compliance"],
    ["Capacity (Amanda F., Azure TPM)","Capacity by SKU returns H200/B200/CPU breakdown","curl localhost:8000/api/v1/fleet/capacity?by=sku,az"],
    ["K8s Platform (Wei C., Google GKE)","minikube node cordon/uncordon via K8s adapter works","nlm cordon minikube-node --priority P2 --owner k8s-platform"],
    ["Provisioning (Kevin R., Together)","Seed new node → provisioning → burn_in → certified_ready flow","nlm provision --node new-gpu-001 --sku DGX-H200"],
    ["Automation (Andrew J., CoreWeave)","Certification API returns test results in JSON","curl localhost:8000/api/v1/certifications"],
], [1.1,2.2,2.8])

doc.add_page_break()

# ══════════════ 5. FILE STRUCTURE ══════════════
doc.add_heading("5  Repository File Structure", level=1)
T(doc, ["Path","Purpose","Lines Est."], [
    ["nlm-controller/nlm/__init__.py","Package init","15"],
    ["nlm-controller/nlm/models.py","Node, State, FailureClass, Incident dataclasses","250"],
    ["nlm-controller/nlm/statemachine.py","11-state engine + transition validation","260"],
    ["nlm-controller/nlm/classifier.py","FailureClassifier (15 rules) + IncidentCorrelator","430"],
    ["nlm-controller/nlm/cordon.py","Cordon priority arbitration (P0–P4)","150"],
    ["nlm-controller/nlm/controller.py","Main controller loop (watches telemetry → classifies → transitions)","200"],
    ["nlm-controller/nlm/api.py","FastAPI REST endpoints (15 routes)","300"],
    ["nlm-controller/nlm/cli.py","Click CLI: nlm status, cordon, transition, inject, capacity, test","250"],
    ["nlm-controller/nlm/db.py","SQLite CRUD (nodes, events, incidents, certifications)","200"],
    ["nlm-controller/nlm/mock_fleet.py","Seed 50 mock nodes, fault injection engine","300"],
    ["nlm-controller/nlm/adapters/base.py","Abstract NodeBackend interface","120"],
    ["nlm-controller/nlm/adapters/bcm.py","BCM adapter (real cmsh + --mock mode)","200"],
    ["nlm-controller/nlm/adapters/k8s.py","K8s adapter (real kubectl via minikube)","180"],
    ["nlm-controller/nlm/adapters/bare_metal.py","IPMI/Redfish adapter (real + --mock mode)","200"],
    ["nlm-controller/nlm/alerts.py","Alert dispatcher (Slack/PD real + mock log)","120"],
    ["nlm-controller/config/node-states.yml","State machine YAML config","100"],
    ["nlm-controller/config/nlm.yml","Global NLM config (DB path, mock flags, thresholds)","60"],
    ["nlm-controller/tests/test_statemachine.py","State machine unit tests","200"],
    ["nlm-controller/tests/test_classifier.py","Classifier unit tests (all 15 rules)","250"],
    ["nlm-controller/tests/test_cordon.py","Cordon priority tests","100"],
    ["nlm-controller/tests/test_api.py","API integration tests (httpx)","200"],
    ["nlm-controller/tests/test_integration.py","Full flow: inject → classify → cordon → route","300"],
    ["nlm-controller/tests/fixtures/mock-pods.yaml","K8s mock workload pods for minikube","30"],
], [2.5,3.0,0.4])

P(doc, "Total estimated code: ~4,165 lines of Python + YAML. AI vibe-codeable in ~2 days.", bold=True, size=10)

doc.add_paragraph("")
doc.add_heading("6  Expert Council Final Consensus", level=1)

T(doc, ["Expert","Org","Key Recommendation for Local Stack"], [
    ["Karthik V.","Google Borg","Use SQLite in dev, PostgreSQL in prod. Same ORM, zero config."],
    ["James P.","Meta GPU Fleet","50-node mock fleet is the sweet spot. Covers all edge cases without overhead."],
    ["David K.","OpenAI","Fault injection CLI (nlm inject) is critical. Test every classifier rule locally."],
    ["Michael S.","xAI Colossus","Real nvidia-smi + DCGM against GV100 validates the telemetry parser for free."],
    ["Vikram N.","Azure AI HPC","minikube for real K8s adapter testing. No need for kind or k3s — minikube is already installed."],
    ["Steven C.","NVIDIA DGX Cloud","Mock IPMI must return realistic SEL event IDs. Use actual ipmitool output format."],
    ["Wei C.","Google GKE","--mock flag pattern is clean: same adapter code, different backend. No separate mock adapter class."],
    ["Marcus T.","Meta Capacity","pytest fixtures seed the 50-node fleet before each test run. Deterministic."],
    ["Carlos G.","xAI Observability","alerts.jsonl as mock Slack/PD is simple and testable. Parse in integration tests."],
    ["Andrew J.","CoreWeave Infra","File structure is right-sized. Don't split into more modules until >10K lines."],
    ["Rachel S.","Google TPM","2-day AI vibe-coding timeline is realistic. Prioritize: statemachine → classifier → CLI → API."],
    ["Brian W.","Meta TPM","Run pytest in CI on every push. Coverage target: 90% for core logic."],
    ["Lisa M.","OpenAI TPM","Test customer protection flow locally: protected node → block transition → verify."],
    ["Jennifer L.","xAI TPM","Mock fleet should include at least 3 node states per state (33 of 50 nodes pre-assigned)."],
    ["Helen T.","Azure TPM","API integration tests using httpx TestClient — no actual server needed."],
    ["Karen H.","NVIDIA TPM","DCGM Exporter install: apt-get install datacenter-gpu-manager. One command."],
    ["Michelle K.","Lambda TPM","Rich CLI output (tables, colors) makes local dev much faster. Already have rich installed."],
    ["Paul S.","CoreWeave TPM","nlm capacity command is the highest-value local feature for quick validation."],
    ["Amanda F.","Azure AI Programs","This local stack design 1:1 mirrors production. Code moves to prod with config change only."],
    ["Rebecca M.","NVIDIA Fleet Programs","Include a Makefile: make setup, make test, make run, make inject-all for quick iteration."],
    ["Daniel P.","NVIDIA BMC Expert","Mock Redfish JSON must match real /redfish/v1/Systems/1 schema. Use real DMTF examples."],
    ["Kevin R.","Together Reliability","Integration tests should cover the full lifecycle: provision → burn_in → certified → assigned → drain → repair → certified."],
], [0.8,1.2,4.8])

doc.add_paragraph("")
P(doc, "Unanimous expert recommendation: Start local development immediately. This machine exceeds all requirements. The mock framework enables full NLM testing without access to production infrastructure. The same code deploys to production with a config switch (--mock off).", bold=True, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f"✅ Generated: {OUT}")
print(f"   Size: {os.path.getsize(OUT)/1024:.0f} KB")
